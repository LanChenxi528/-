# coding:utf-8

# from langchain.agents import AgentExecutor
# # from search import DeepSearch
# from tool import Search_www_Tool
# from intent_agent import IntentAgent
# # from llm_model import  ChatGLM
# from model_loader import get_loaded_model
from docx import Document
from docx.shared import Pt
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from datetime import datetime


import os
# end="""标题：2021年4月2日台湾铁路太鲁阁号脱轨事故简报
#
# 一、事故基本情况
# 2021年4月2日9时许，台湾铁路“太鲁阁号”408次列车在花莲大清水隧道内脱轨，事故是由于被滑落铁轨的工程车砸中导致。事故造成48人死亡，118人受伤。列车为树林往台东线路，事故发生时，列车时速约120公里/时。台铁“太鲁阁号”408次列车共有8节车厢，载有约500名乘客。事故导致第2、第3车厢出轨，第3至第8车厢位于隧道内且部分变形。
#
# 二、应急处置情况
# 事故发生后，现场紧急疏散乘客，消防和警察全力抢救。根据初步研判，事故原因是隧道上方施工厂商未将工程车停稳，导致其滑落撞击列车。台铁部门表示将尽快抢修事故路段。受伤人员被分别送至当地多家医院救治。对于事故原因，涉嫌肇事的相关负责人已被警方带回侦讯。现场救援队伍组织救援工作，同时官方负责人对事故表示哀悼，并承诺进行彻底调查及采取补救措施。由于事故影响，相关线路被迫关闭，并对列车运行时间进行了调整。
#
# 三、相关事故列表
# 2004年5月16日，美国发生铁路事故，事故原因为铁路公司游说导致的安全规定不力，最终以提高铁路安全的新规被推迟和修改告终。
# 2022年2月7日，美国佐治亚州道格拉斯发生列车出轨事故，事故原因是卡车陷在轨道之上，列车与之相撞而脱轨，造成列车人员受伤。同时，2022年2月14日，德国慕尼黑附近的两列通勤列车迎面相撞，造成1人死亡，至少14人受伤。
# 2022年3月8日21时28分，甘肃省白银市靖远县境内红会线发生一起重型自卸货车碰撞铁路桥梁致正常通过的31024次货物列车脱轨铁路交通较大事故。事故调查结果已公布，事故原因为货车超载行驶，事故处置方式为现场救援并依法处理。
# 2023年4月23日，美国俄亥俄州东巴勒斯坦城，由于一列运有危险化学品的火车脱轨，引发化学品泄漏事故，事故原因尚在调查中，目前已有超过1000列火车脱轨，事故处置方式正在积极进行中。
# 2023年2月24日，美国一有毒化学品运载列车脱轨事故初步调查报告发布，尚未确定事故原因，两党将焦点放在政治角力方面。
#
# 参考站源：
# https://www.thepaper.cn/newsDetail_forward_12018460
# http://www.xinhuanet.com/2021-04/02/c_1127289601.htm"""

# 创建一个新的word 文档
# doc = Document()

# 将正文内容分割为标题和段落
# lines = end.split('\n', 1)
# title = lines[0]  # 第一行作为标题
# title1 = title[3:]
# if len(lines) > 1:
#     # 剩余的文本作为段落内容paragraph content = lines[1]else:
#     paragraph_content = lines[1]
# else:
#     paragraph_content = ''
def Doc_Generate(id,end):
    # 使用正则表达式分割文本
    parts = re.split('一、事故基本情况|二、应急处置情况|三、相关事故列表|参考站源：', end)
    # 标题
    title = parts[0].strip().replace('标题：','')
    print("模块1的word文档的标题是："+title)
    # title = parts[1].strip()
    # 事故基本情况
    accident_info = parts[1].strip()
    # print(accident_info)
    # 应急处置情况
    emergency_response = parts[2].strip()
    # print(emergency_response)
    # 相关事故列表
    related_list = parts[3].strip()
    # print(related_list)
    # 参考站源
    references = parts[4].strip()
    # print(references)

    Doc = Document()
    Doc.styles['Normal'].font.name = u'Times New Roman'
    Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    Doc.styles['Normal'].font.size = Pt(16)
    Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    Doc.styles['Normal'].paragraph_format.line_spacing = Pt(28)

    #标题部分
    # Head = Doc.add_heading("",level=1)# 这里不填标题内容
    # # 设置标题段落的对齐方式为居中
    # Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # run  = Head.add_run(title+"\n")
    Head = Doc.add_paragraph()#Head = Doc.add_heading("", level=1)  # 这里不填标题内容
    # 设置标题段落的对齐方式为居中
    Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = Head.add_run(title + "\n")
    run.bold = False
    run.font.size=Pt(22)
    run.font.name=u'Times New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')
    #一、事故基本情况部分
    Head1 = Doc.add_paragraph()
    Head1.paragraph_format.first_line_indent = Pt(32)
    run1  = Head1.add_run("一、事故基本情况")
    run1.bold = False
    run1.font.size=Pt(16)
    run1.font.name=u'Times New Roman'
    run1.font.color.rgb = RGBColor(0,0,0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # Doc.add_paragraph(accident_info)
    # 添加带有首行缩进的段落，并设置文本内容
    paragraph = Doc.add_paragraph()
    run = paragraph.add_run(accident_info)
    paragraph.paragraph_format.first_line_indent = Pt(32)
    #二、应急处置情况部分
    Head1 = Doc.add_paragraph()#Head1 = Doc.add_heading("",level=2)# 这里不填标题内容
    Head1.paragraph_format.first_line_indent = Pt(32)
    run1  = Head1.add_run("二、应急处置情况")
    run1.bold = False
    run1.font.size=Pt(16)
    run1.font.name=u'Times New Roman'
    run1.font.color.rgb = RGBColor(0,0,0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # Doc.add_paragraph(emergency_response)
    paragraph = Doc.add_paragraph()
    run = paragraph.add_run(emergency_response)
    paragraph.paragraph_format.first_line_indent = Pt(32)
    #三、相关事故列表部分
    Head1 = Doc.add_paragraph()#Head1 = Doc.add_heading("",level=2)# 这里不填标题内容
    Head1.paragraph_format.first_line_indent = Pt(32)
    run1  = Head1.add_run("三、历史上的类似事故")
    run1.bold = False
    run1.font.size=Pt(16)
    run1.font.name=u'Times New Roman'
    run1.font.color.rgb = RGBColor(0,0,0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # Doc.add_paragraph(related_list)

    accident_dates = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', related_list)
    print("看看模块1的word中历史事故列表都有什么\n")
    print(accident_dates)
    # related_events = []
    for i, date in enumerate(accident_dates):
        start_index = related_list.find(date)
        if i < len(accident_dates) - 1:
            next_date_index = related_list.find(accident_dates[i + 1])
            if next_date_index != -1:
                end_index = next_date_index
            else:
                end_index = len(related_list)
        else:
            end_index = len(related_list)
        event = related_list[start_index:end_index].strip()
        # Doc.add_paragraph(event)
        paragraph = Doc.add_paragraph()
        run = paragraph.add_run(event)
        paragraph.paragraph_format.first_line_indent = Pt(32)
    # for paragraph in Doc.paragraphs:
    #     # 计算一个字符的宽度（大约12磅左右）
    #     char_width = 12
    #     # 设置首行缩进为2个字符的宽度
    #     first_line_indent = 2 * char_width
    #     paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)

    # 参考站源部分
    Head1 = Doc.add_paragraph()#heading("",level=2)# 这里不填标题内容
    # Head1.paragraph_format.first_line_indent = Pt(24)
    run1  = Head1.add_run("\n参考站源：")
    run1.bold = False
    run1.font.size=Pt(14)
    run1.font.name=u'Times New Roman'
    run1.font.color.rgb = RGBColor(0,0,0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    references_urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', references)
    rus=[]
    for i, date in enumerate(references_urls):
        paragraph = Doc.add_paragraph()
        run = paragraph.add_run(date)
        run.font.size = Pt(14)
        run.font.name = u'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')

    paragraph = Doc.add_paragraph()
    run = paragraph.add_run("\n\n铁科智问大模型")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 获取当前日期和时间
    now = datetime.now()
    # 获取年、月、日
    year = now.year
    month = now.month
    day = now.day
    string1=str(year)+'年'+str(month)+"月"+str(day)+"日"
    paragraph2 = Doc.add_paragraph()
    run2 = paragraph2.add_run(string1)
    run2.font.name=u'Times New Roman'
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # print("年份:", year)
    # print("月份:", month)
    # print("日期:", day)

    # Doc.add_paragraph("Python")
    # Doc.add_paragraph("Python 对word进行操作")

    word_filename = id  + '.docx'
    word_file = "word_save/"
    word_filepath = os.path.join(word_file, word_filename)
    Doc.save(word_filepath)
    # Doc.save("Python_word.docx")

# # 如果有剩余的文本，将其添加为段落if paragraph content:
# if paragraph_content:
#     paragraph = doc.add_paragraph(paragraph_content)

# # 保存文档
# word_filename = title1 + '.docx'
# word_file = "../word_save/"
# word_filepath = os.path.join(word_file, word_filename)
#
# doc.save(word_filepath)
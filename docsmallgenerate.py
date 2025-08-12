from docx import Document
from docx.shared import Pt
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from datetime import datetime


import os
def Doc_SmallGenerate(id1,end):
    parts = re.split('标题：|一、事故概况：|参考站源：', end)
    # 标题
    title = parts[1].strip().replace('标题:', '')
    print(title)
    # title = parts[1].strip()
    # 一、事故概况：
    related_list  = parts[2].strip()
    print(related_list )

    references = parts[3].strip()
    print(references)

    Doc1 = Document()
    Doc1.styles['Normal'].font.name = u'Times New Roman'
    Doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    Doc1.styles['Normal'].font.size = Pt(16)
    Doc1.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    Doc1.styles['Normal'].paragraph_format.line_spacing = Pt(28)

    # 标题部分
    # Head = Doc.add_heading("",level=1)# 这里不填标题内容
    # # 设置标题段落的对齐方式为居中
    # Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # run  = Head.add_run(title+"\n")
    Head = Doc1.add_paragraph()  # Head = Doc.add_heading("", level=1)  # 这里不填标题内容
    # 设置标题段落的对齐方式为居中
    Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = Head.add_run(title)
    run.bold = False
    run.font.size = Pt(22)
    run.font.name = u'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')
    # 事故概况
    Head1 = Doc1.add_paragraph()  # Head1 = Doc.add_heading("",level=2)# 这里不填标题内容
    Head1.paragraph_format.first_line_indent = Pt(32)
    run1 = Head1.add_run("一、事故概况")
    run1.bold = False
    run1.font.size = Pt(16)
    run1.font.name = u'Times New Roman'
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    accident_dates = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', related_list)
    for i, date in enumerate(accident_dates):
        # start_index = related_list.find(date)
        if i==0:start_index = related_list.find(date)
        if i < len(accident_dates) - 1:
            next_date_index = related_list.find(accident_dates[i + 1], start_index + 1)
            if next_date_index != -1:
                end_index = next_date_index
            else:
                end_index = len(related_list)
        else:
            end_index = len(related_list)
        event = related_list[start_index:end_index].strip()
        start_index=end_index
        # Doc.add_paragraph(event)
        paragraph = Doc1.add_paragraph()
        run = paragraph.add_run(event)
        paragraph.paragraph_format.first_line_indent = Pt(32)

    Head1 = Doc1.add_paragraph()
    run1 = Head1.add_run("\n参考站源：")
    run1.bold = False
    run1.font.size = Pt(14)
    run1.font.name = u'Times New Roman'
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    references_urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
                                 references)

    for i, date in enumerate(references_urls):
        paragraph = Doc1.add_paragraph()
        run1 = paragraph.add_run(date)
        run1.font.size = Pt(14)
        run1.font.name = u'Calibri'
        run1.font.color.rgb = RGBColor(0, 0, 0)
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')

    paragraph = Doc1.add_paragraph()
    run = paragraph.add_run("\n\n铁科智问大模型")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 获取当前日期和时间
    now = datetime.now()
    # 获取年、月、日
    year = now.year
    month = now.month
    day = now.day
    string1 = str(year) + '年' + str(month) + "月" + str(day) + "日"
    paragraph2 = Doc1.add_paragraph()
    run2 = paragraph2.add_run(string1)
    run2.font.name = u'Times New Roman'
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # print("年份:", year)
    # print("月份:", month)
    # print("日期:", day)

    # Doc.add_paragraph("Python")
    # Doc.add_paragraph("Python 对word进行操作")
    word_filename1 = id1 + '.docx'
    word_file1 = "word_save/"
    word_filepath1 = os.path.join(word_file1, word_filename1)
    Doc1.save(word_filepath1)